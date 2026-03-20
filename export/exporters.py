import html
import re
import logging
from pathlib import Path
from typing import List, Dict, Protocol
from datetime import datetime
from dataclasses import dataclass, field

logger = logging.getLogger(__name__)

def _validate_hex_color(color: str) -> str:
    if color and re.fullmatch(r'#[0-9a-fA-F]{3,6}', color.strip()):
        return color.strip()
    return '#4f46e5'

@dataclass
class ExportMetadata:
    code_name: str
    code_color: str
    segments: List[Dict]
    export_date: datetime = field(default_factory=datetime.now)

class ExportFormatter(Protocol):
    def format(self, metadata: ExportMetadata) -> str:
        ...

class TextFormatter:
    def format(self, metadata: ExportMetadata) -> str:
        lines = []
        lines.append(f"╔{'═' * 58}╗")
        lines.append(f"║  KOD RAPORU: {metadata.code_name.upper():<44}║")
        lines.append(f"╠{'═' * 58}╣")
        lines.append(f"║  Tarih: {metadata.export_date.strftime('%d.%m.%Y %H:%M'):<48}║")
        lines.append(f"║  Toplam Segment: {len(metadata.segments):<40}║")
        lines.append(f"╚{'═' * 58}╝\n")
        
        current_doc = None
        for i, seg in enumerate(metadata.segments, 1):
            doc_title = seg.get('document_title', 'Bilinmeyen')
            if doc_title != current_doc:
                current_doc = doc_title
                lines.append(f"\n{'─' * 60}")
                lines.append(f"📄 {current_doc}")
                lines.append(f"{'─' * 60}\n")
            
            lines.append(f"[{i}] \"{seg.get('segment_text', '')}\"")
            lines.append(f"    └─ Konum: karakter {seg.get('start_pos', 0)}-{seg.get('end_pos', 0)}\n")
            
        return "\n".join(lines)

class MarkdownFormatter:
    def format(self, metadata: ExportMetadata) -> str:
        lines = [f"# Kod Raporu: {metadata.code_name}\n"]
        lines.append(f"> Oluşturulma: {metadata.export_date.strftime('%d.%m.%Y %H:%M')} | Toplam: {len(metadata.segments)} segment\n")
        
        current_doc = None
        for i, seg in enumerate(metadata.segments, 1):
            doc_title = seg.get('document_title', 'Bilinmeyen')
            if doc_title != current_doc:
                current_doc = doc_title
                lines.append(f"## 📄 {doc_title}\n")
            
            lines.append(f"### Segment {i}")
            lines.append(f"> {seg.get('segment_text', '')}\n")
            lines.append(f"- **Konum:** {seg.get('start_pos', 0)}-{seg.get('end_pos', 0)}")
            if seg.get('weight'):
                lines.append(f"- **Ağırlık:** {'⭐' * seg.get('weight')}")
            lines.append("\n---\n")
        return "\n".join(lines)

class HTMLFormatter:
    def format(self, metadata: ExportMetadata) -> str:
        safe_color = _validate_hex_color(metadata.code_color)
        html_out = f"""<!DOCTYPE html>
<html lang="tr">
<head>
    <meta charset="UTF-8">
    <title>Kod Raporu: {html.escape(metadata.code_name)}</title>
    <style>
        body {{ font-family: sans-serif; max-width: 800px; margin: 40px auto; background: #f8fafc; color: #1e293b; padding: 20px; }}
        .header {{ background: white; padding: 30px; border-radius: 12px; border-left: 6px solid {safe_color}; box-shadow: 0 4px 6px rgba(0,0,0,0.05); text-align: center; margin-bottom: 40px; }}
        h1 {{ color: {safe_color}; }}
        .segment {{ background: white; padding: 20px; margin: 15px 0; border-radius: 8px; border-left: 4px solid {safe_color}; box-shadow: 0 2px 4px rgba(0,0,0,0.05); }}
        .document-title {{ color: #4f46e5; border-bottom: 2px solid #e2e8f0; padding: 10px 0; }}
    </style>
</head>
<body>
    <div class="header">
        <h1>🏷️ {html.escape(metadata.code_name)}</h1>
        <p>Tarih: {metadata.export_date.strftime('%d.%m.%Y %H:%M')} | Toplam: {len(metadata.segments)}</p>
    </div>"""
        
        current_doc = None
        for i, seg in enumerate(metadata.segments, 1):
            doc_title = html.escape(seg.get('document_title', 'Bilinmeyen'))
            if doc_title != current_doc:
                current_doc = doc_title
                html_out += f'<h2 class="document-title">📄 {doc_title}</h2>'
            
            txt = html.escape(seg.get('segment_text', ''))
            html_out += f"""<div class="segment">
                <strong>[{i}]</strong> <em>"{txt}"</em>
                <br><small style="color:#94a3b8">Konum: {seg.get('start_pos')}-{seg.get('end_pos')}</small>
            </div>"""
            
        html_out += "</body></html>"
        return html_out

_INJECTION_PREFIXES = ("=", "+", "-", "@", "\t")

def _sanitize_spreadsheet_cell(value):
    if value is None:
        return ""
    text = str(value)
    if text and text[0] in _INJECTION_PREFIXES:
        return f"'{text}"
    return text

def export_segments(file_path: str, formatter: ExportFormatter, metadata: ExportMetadata) -> bool:
    """Unified entry point for exports using Strategy pattern."""
    try:
        content = formatter.format(metadata)
        with open(file_path, 'w', encoding='utf-8') as f:
            f.write(content)
        return True
    except Exception as e:
        logger.error(f"Export failed: {e}", exc_info=True)
        return False

def export_to_txt(file_path: str, code_name: str, code_color: str, segments: List[Dict]) -> bool:
    meta = ExportMetadata(code_name, code_color, segments)
    return export_segments(file_path, TextFormatter(), meta)

def export_to_docx(file_path: str, code_name: str, code_color: str, segments: List[Dict]) -> bool:
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        
        doc = Document()
        title = doc.add_heading(f'Kod Raporu: {code_name}', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        meta_p = doc.add_paragraph()
        meta_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        meta_p.add_run(f"Tarih: {datetime.now().strftime('%d.%m.%Y %H:%M')}\n").italic = True
        meta_p.add_run(f"Toplam Segment: {len(segments)}").italic = True
        
        current_doc = None
        for i, seg in enumerate(segments, 1):
            doc_title = seg.get('document_title', 'Bilinmeyen')
            if doc_title != current_doc:
                current_doc = doc_title
                doc.add_heading(f'📄 {doc_title}', level=1)
            
            p = doc.add_paragraph()
            p.add_run(f'[{i}] ').bold = True
            p.add_run(f'"{seg.get("segment_text", "")}"').italic = True
            
            pos_p = doc.add_paragraph()
            pos_run = pos_p.add_run(f'    Konum: {seg.get("start_pos")}-{seg.get("end_pos")}')
            pos_run.font.size = Pt(9)
            pos_run.font.color.rgb = RGBColor(100, 116, 139)
        
        doc.save(file_path)
        return True
    except Exception as e:
        logger.error(f"DOCX export error: {e}", exc_info=True)
        return False

def export_memos_to_docx(file_path: str, memos: List[Dict]) -> bool:
    """Export all project memos to a formatted Word document."""
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        
        doc = Document()
        title = doc.add_heading('PROJE MEMO RAPORU', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        meta_p = doc.add_paragraph()
        meta_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        meta_p.add_run(f"Oluşturulma: {datetime.now().strftime('%d.%m.%Y %H:%M')}\n").italic = True
        meta_p.add_run(f"Toplam Not: {len(memos)}").italic = True
        
        # Group memos by category
        segment_memos = [m for m in memos if m.get('start_pos')]
        code_memos = [m for m in memos if m.get('code_id') and not m.get('start_pos')]
        doc_memos = [m for m in memos if m.get('document_id') and not m.get('start_pos') and not m.get('code_id')]
        
        if doc_memos:
            doc.add_heading('📂 Belge Notları', level=1)
            for m in doc_memos:
                doc.add_heading(f"📄 {m.get('doc_title', 'Belge')}", level=2)
                p = doc.add_paragraph(m.get('content', ''))
                meta = doc.add_paragraph()
                meta.add_run(f"Tarih: {m.get('created_at')}").italic = True
                meta.style.font.size = Pt(8)
                
        if code_memos:
            doc.add_heading('🏷️ Kod Notları', level=1)
            for m in code_memos:
                doc.add_heading(f"Kod: {m.get('title', 'İsimsiz Kod')}", level=2)
                p = doc.add_paragraph(m.get('content', ''))
                
        if segment_memos:
            doc.add_heading('📑 Segment/Alıntı Notları', level=1)
            for m in segment_memos:
                doc.add_heading(f"Bağlam: {m.get('doc_title', 'Belge')} ({m.get('start_pos')}-{m.get('end_pos')})", level=2)
                p = doc.add_paragraph(m.get('content', ''))
        
        doc.save(file_path)
        return True
    except Exception as e:
        logger.error(f"Memo DOCX export error: {e}", exc_info=True)
        return False

def export_to_html(file_path: str, code_name: str, code_color: str, segments: List[Dict]) -> bool:
    meta = ExportMetadata(code_name, code_color, segments)
    return export_segments(file_path, HTMLFormatter(), meta)

def export_to_csv(file_path: str, code_name: str, code_color: str, segments: List[Dict]) -> bool:
    try:
        import pandas as pd
        df = pd.DataFrame(segments)
        cols = ['document_title', 'segment_text', 'start_pos', 'end_pos', 'weight', 'created_at']
        available_cols = [c for c in cols if c in df.columns]
        df = df[available_cols]
        for col in ('document_title', 'segment_text'):
            if col in df.columns:
                df[col] = df[col].map(_sanitize_spreadsheet_cell)
        df.to_csv(file_path, index=False, encoding='utf-8-sig')
        return True
    except Exception as e:
        logger.error(f"CSV export error: {e}", exc_info=True)
        return False

def export_to_xlsx(file_path: str, code_name: str, code_color: str, segments: List[Dict]) -> bool:
    try:
        import pandas as pd
        df = pd.DataFrame(segments)
        cols = ['document_title', 'segment_text', 'start_pos', 'end_pos', 'weight', 'created_at']
        available_cols = [c for c in cols if c in df.columns]
        df = df[available_cols]
        for col in ('document_title', 'segment_text'):
            if col in df.columns:
                df[col] = df[col].map(_sanitize_spreadsheet_cell)
        # Rename columns for better Excel display
        rename_map = {
            'document_title': 'Belge Adı',
            'segment_text': 'Kodlu Bölüm (Segment)',
            'start_pos': 'Başlangıç',
            'end_pos': 'Bitiş',
            'weight': 'Ağırlık',
            'created_at': 'Tarih'
        }
        df = df.rename(columns=rename_map)
        df.to_excel(file_path, index=False)
        return True
    except Exception as e:
        logger.error(f"XLSX export error: {e}", exc_info=True)
        return False

def export_to_json(file_path: str, code_name: str, code_color: str, segments: List[Dict]) -> bool:
    try:
        import json
        data = {
            "code_name": code_name,
            "export_date": datetime.now().isoformat(),
            "total_segments": len(segments),
            "segments": segments
        }
        with open(file_path, 'w', encoding='utf-8') as f:
            json.dump(data, f, ensure_ascii=False, indent=4)
        return True
    except Exception as e:
        logger.error(f"JSON export error: {e}", exc_info=True)
        return False

def export_to_markdown(file_path: str, code_name: str, code_color: str, segments: List[Dict]) -> bool:
    meta = ExportMetadata(code_name, code_color, segments)
    return export_segments(file_path, MarkdownFormatter(), meta)

def get_export_formats() -> List[Dict]:
    formats = [
        {'name': 'Metin Dosyası', 'extension': 'txt', 'filter': 'Metin Dosyaları (*.txt)', 'function': export_to_txt},
        {'name': 'HTML', 'extension': 'html', 'filter': 'HTML Dosyaları (*.html)', 'function': export_to_html},
        {'name': 'CSV', 'extension': 'csv', 'filter': 'CSV Dosyaları (*.csv)', 'function': export_to_csv},
        {'name': 'Excel Belgesi', 'extension': 'xlsx', 'filter': 'Excel Dosyaları (*.xlsx)', 'function': export_to_xlsx},
        {'name': 'Markdown', 'extension': 'md', 'filter': 'Markdown Dosyaları (*.md)', 'function': export_to_markdown},
        {'name': 'JSON', 'extension': 'json', 'filter': 'JSON Dosyaları (*.json)', 'function': export_to_json}
    ]
    try:
        import docx
        formats.insert(1, {'name': 'Word Belgesi', 'extension': 'docx', 'filter': 'Word Belgeleri (*.docx)', 'function': export_to_docx})
    except ImportError:
        pass
    return formats
