"""
NER Export Module
Provides Word and HTML export for named entity recognition results.
"""

from datetime import datetime
from typing import Dict


def export_ner_to_word(file_path: str, ner_data: Dict, model_type: str = "NER") -> bool:
    try:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        doc = Document()
        title = doc.add_heading('Varlık Tanıma (NER) Raporu', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        meta = doc.add_paragraph()
        meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
        meta.add_run(f'Model: {model_type}\n').italic = True
        meta.add_run(f'Tarih: {datetime.now().strftime("%d.%m.%Y %H:%M")}\n').italic = True
        meta.add_run(f'Toplam Belge: {len(ner_data.get("documents", []))}\n').italic = True
        meta.add_run(f'Toplam Varlık: {sum(ner_data.get("summary", {}).values())}\n').italic = True

        doc.add_paragraph()
        doc.add_heading('1. Varlık Kategorileri', level=1)
        all_entities = ner_data.get("all_entities", {})
        label_names = {
            "PER": "Kişi",
            "LOC": "Yer",
            "ORG": "Kurum",
            "DATE": "Tarih",
            "MISC": "Diğer"
        }
        for label in ["PER", "LOC", "ORG", "DATE", "MISC"]:
            entities = sorted(set(all_entities.get(label, [])))
            if not entities:
                continue
            doc.add_heading(label_names.get(label, label), level=2)
            doc.add_paragraph(", ".join(entities))

        documents = ner_data.get("documents", [])
        if ner_data.get("mode") == "hybrid":
            doc.add_heading('2. Belge Bazlı Hibrit Karşılaştırma', level=1)
            for item in documents:
                doc.add_heading(item.get("title", "Belge"), level=2)
                local_entities = ", ".join(ent.get("text", "") for ent in item.get("local_entities", [])) or "-"
                online_entities = ", ".join(ent.get("text", "") for ent in item.get("online_entities", [])) or "-"
                comp = item.get("comparison", {})
                doc.add_paragraph(f"Lokal: {local_entities}")
                doc.add_paragraph(f"AI: {online_entities}")
                doc.add_paragraph(f"Ortak: {comp.get('shared_count', 0)} | Sadece Lokal: {comp.get('local_only_count', 0)} | Sadece AI: {comp.get('online_only_count', 0)}")
        else:
            doc.add_heading('2. Belge Bazlı Dağılım', level=1)
            table = doc.add_table(rows=1, cols=3)
            table.style = 'Light Grid Accent 1'
            hdr = table.rows[0].cells
            hdr[0].text = 'Belge'
            hdr[1].text = 'Sayı'
            hdr[2].text = 'Varlıklar'
            for item in documents:
                entities = item.get("entities", [])
                row = table.add_row().cells
                row[0].text = item.get("title", "Belge")
                row[1].text = str(len(entities))
                row[2].text = ", ".join(ent.get("text", "") for ent in entities)

        footer = doc.add_paragraph()
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer.add_run('LexiScholar Akademik Analiz Yazılımı').italic = True
        doc.save(file_path)
        return True
    except Exception as e:
        print(f"NER Word export error: {e}")
        return False


def export_ner_to_html(file_path: str, ner_data: Dict, model_type: str = "NER") -> bool:
    try:
        mode = ner_data.get("mode", "local")
        if mode == "hybrid":
            from ui.visualizations.semantic_analytics import generate_hybrid_entities_html
            generated_path = generate_hybrid_entities_html(ner_data, model_name=model_type)
        elif mode == "online":
            from ui.visualizations.semantic_analytics import generate_online_entities_html
            generated_path = generate_online_entities_html(ner_data, model_name=model_type)
        else:
            from ui.visualizations.semantic_analytics import generate_entities_html
            generated_path = generate_entities_html(ner_data)

        with open(generated_path, 'r', encoding='utf-8') as src:
            html_content = src.read()
        with open(file_path, 'w', encoding='utf-8') as dst:
            dst.write(html_content)
        return True
    except Exception as e:
        print(f"NER HTML export error: {e}")
        return False
