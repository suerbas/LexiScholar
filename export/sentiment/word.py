"""
Export Sentiment Analysis to Word.
"""

from datetime import datetime
from typing import List, Dict, Optional
from .utils import _translate_label

def export_sentiment_to_word(
    file_path: str,
    results: List[Dict],
    model_type: str = "BERT",
    hybrid_results: Optional[List[Dict]] = None
) -> bool:
    """
    Export sentiment analysis results to Word (.docx) format.
    """
    try:
        from docx import Document
        from docx.shared import Inches, Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_TABLE_ALIGNMENT
        
        doc = Document()
        
        # Title
        title = doc.add_heading('Duygu Analizi Raporu', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Metadata
        meta = doc.add_paragraph()
        meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
        meta.add_run(f'\nModel: {model_type}\n').italic = True
        meta.add_run(f'Tarih: {datetime.now().strftime("%d.%m.%Y %H:%M")}\n').italic = True
        meta.add_run(f'Toplam Belge: {len(results)}\n').italic = True
        
        doc.add_paragraph()
        
        # Summary statistics
        doc.add_heading('1. Özet İstatistikler', level=1)
        _add_summary_stats_to_doc(doc, results)
        
        # Detailed results
        doc.add_heading('2. Belge Bazlı Sonuçlar', level=1)
        
        for i, r in enumerate(results, 1):
            # Document heading
            p = doc.add_paragraph()
            p.add_run(f'{i}. {r.get("title", "Bilinmeyen")}').bold = True
            
            # Sentiment info
            label = _translate_label(r.get('label', 'neutral'))
            score = r.get('score', 0.5)
            summary = r.get('summary', '')
            
            info = doc.add_paragraph()
            info.add_run(f'Duygu: {label} | Güven: {score:.1%}')
            if summary:
                info.add_run(f' | {summary}')
            
            doc.add_paragraph()  # Spacing
        
        # Hybrid comparison if available
        if hybrid_results:
            doc.add_heading('3. Model Karşılaştırması', level=1)
            _add_hybrid_comparison_to_doc(doc, hybrid_results)
        
        # Footer
        doc.add_paragraph()
        footer = doc.add_paragraph()
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer.add_run('LexiScholar Akademik Analiz Yazılımı').italic = True
        
        doc.save(file_path)
        return True
        
    except Exception as e:
        print(f"Word export error: {e}")
        return False

def export_hybrid_sentiment_to_word(file_path: str, results: List[Dict], model_type: str = "AI") -> bool:
    """
    Export hybrid sentiment comparison to Word.
    """
    try:
        from docx import Document
        from docx.shared import Inches, Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        
        doc = Document()
        
        # Title
        title = doc.add_heading('Hibrit Duygu Analizi Raporu', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Metadata
        meta = doc.add_paragraph()
        meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
        meta.add_run(f'Model: BERT + {model_type} | ').italic = True
        meta.add_run(f'Tarih: {datetime.now().strftime("%d.%m.%Y %H:%M")} | ').italic = True
        meta.add_run(f'Toplam: {len(results)} belge').italic = True
        
        doc.add_paragraph()
        
        # Calculate statistics
        match_count = 0
        for r in results:
            local = r.get('local', {})
            online = r.get('online', {})
            if local.get('label') == online.get('label'):
                match_count += 1
        
        match_rate = f"{(match_count/len(results)*100):.1f}%" if results else "0%"
        
        # Summary statistics
        stats_heading = doc.add_heading('1. Model Uyum İstatistikleri', level=1)
        
        stats = doc.add_paragraph()
        stats.add_run(f'Eşleşme Oranı: ').bold = True
        stats.add_run(f'{match_rate} ({match_count}/{len(results)} belge)')
        
        doc.add_paragraph()
        
        # Detailed comparison table
        doc.add_heading('2. Belge Bazlı Karşılaştırma', level=1)
        
        for i, r in enumerate(results, 1):
            local = r.get('local', {})
            online = r.get('online', {})
            
            doc_title = r.get('title', f'Belge {i}')
            
            # Document heading
            doc.add_heading(f'{i}. {doc_title}', level=2)
            
            # Create comparison table for this document
            table = doc.add_table(rows=3, cols=3)
            table.style = 'Table Grid'
            
            # Header row
            header_cells = table.rows[0].cells
            header_cells[0].text = 'Model'
            header_cells[1].text = 'Duygu'
            header_cells[2].text = 'Skor / Güven'
            
            # Make header bold
            for cell in header_cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True
            
            # BERT row
            bert_cells = table.rows[1].cells
            bert_cells[0].text = '🤖 BERT (Lokal)'
            l_label = _translate_label(local.get('label', 'neutral'))
            l_score = local.get('score', 0.5)
            bert_cells[1].text = l_label
            bert_cells[2].text = f'{l_score:.1%}'
            
            # AI row
            ai_cells = table.rows[2].cells
            ai_cells[0].text = f'🤖 {model_type} (Online)'
            o_label = _translate_label(online.get('label', 'neutral'))
            o_score = online.get('score', 0.5)
            o_conf = online.get('confidence', 0.5)
            ai_cells[1].text = o_label
            ai_cells[2].text = f'Skor: {o_score:.1%} | Güven: {o_conf:.1%}'
            
            # Color code the cells based on match
            is_match = local.get('label') == online.get('label')
            
            # Add summary
            doc.add_paragraph()
            summary_para = doc.add_paragraph()
            summary_para.add_run('BERT Özeti: ').bold = True
            summary_para.add_run(local.get('summary', 'Özet yok'))
            
            summary_para2 = doc.add_paragraph()
            summary_para2.add_run(f'{model_type} Özeti: ').bold = True
            summary_para2.add_run(online.get('summary', 'Özet yok'))
            
            # Match indicator
            match_para = doc.add_paragraph()
            if is_match:
                match_run = match_para.add_run('✓ Modeller uyumlu')
                match_run.font.color.rgb = RGBColor(0x10, 0xB9, 0x81)
                match_run.bold = True
            else:
                match_run = match_para.add_run('⚠ Modeller farklı sonuç verdi')
                match_run.font.color.rgb = RGBColor(0xF5, 0x9E, 0x0B)
                match_run.bold = True
            
            doc.add_paragraph()
        
        # Footer
        doc.add_paragraph()
        footer = doc.add_paragraph()
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer.add_run('LexiScholar Akademik Analiz Yazılımı').italic = True
        
        doc.save(file_path)
        return True
        
    except Exception as e:
        print(f"Hybrid Word export error: {e}")
        return False

def _add_summary_stats_to_doc(doc, results):
    """Add summary statistics to Word document."""
    labels = [r.get('label', 'neutral') for r in results]
    label_counts = {}
    for l in labels:
        label_counts[l] = label_counts.get(l, 0) + 1
    
    # Create summary table
    table = doc.add_table(rows=len(label_counts) + 1, cols=2)
    table.style = 'Light Grid Accent 1'
    
    # Header
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Duygu'
    hdr_cells[1].text = 'Sayı'
    
    # Data
    for i, (label, count) in enumerate(sorted(label_counts.items()), 1):
        row_cells = table.rows[i].cells
        row_cells[0].text = _translate_label(label)
        row_cells[1].text = str(count)

def _add_hybrid_comparison_to_doc(doc, hybrid_results):
    """Add hybrid model comparison to Word document."""
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Light Grid Accent 1'
    
    # Header
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Belge'
    hdr_cells[1].text = 'BERT'
    hdr_cells[2].text = 'Online AI'
    hdr_cells[3].text = 'Uyum'
    
    # Data
    for r in hybrid_results:
        local = r.get('local', {})
        online = r.get('online', {})
        is_match = local.get('label') == online.get('label')
        
        row_cells = table.add_row().cells
        row_cells[0].text = r.get('title', 'Bilinmeyen')
        row_cells[1].text = _translate_label(local.get('label', 'neutral'))
        row_cells[2].text = _translate_label(online.get('label', 'neutral'))
        row_cells[3].text = '✓' if is_match else '✗'
