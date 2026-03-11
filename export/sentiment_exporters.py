"""
Sentiment Analysis Export Module
Provides Excel and Word export for sentiment analysis results.
"""

import os
import re
from datetime import datetime
from typing import List, Dict, Optional
from pathlib import Path

def export_sentiment_to_excel(
    file_path: str,
    results: List[Dict],
    model_type: str = "BERT",
    include_stats: bool = True
) -> bool:
    """
    Export sentiment analysis results to Excel (.xlsx) format.
    
    Args:
        file_path: Target file path
        results: List of sentiment result dictionaries
        model_type: "BERT" or "Online AI"
        include_stats: Include summary statistics sheet
    """
    try:
        import pandas as pd
        from openpyxl import Workbook
        from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
        from openpyxl.utils.dataframe import dataframe_to_rows
        
        # Prepare data
        data = []
        for r in results:
            data.append({
                'Belge Adı': r.get('title', 'Bilinmeyen'),
                'Duygu Etiketi': _translate_label(r.get('label', 'neutral')),
                'Güven Skoru': f"{r.get('score', 0.5):.1%}",
                'Düz Skor': r.get('score', 0.5),
                'Seviye': r.get('level', 3),
                'Özet': r.get('summary', ''),
                'Model': model_type,
                'Tarih': datetime.now().strftime('%d.%m.%Y %H:%M')
            })
        
        df = pd.DataFrame(data)
        
        # Create Excel with formatting
        wb = Workbook()
        ws = wb.active
        ws.title = "Duygu Analizi"
        
        # Write headers
        headers = list(df.columns)
        for col_idx, header in enumerate(headers, 1):
            cell = ws.cell(row=1, column=col_idx, value=header)
            cell.font = Font(bold=True, color="FFFFFF")
            cell.fill = PatternFill(start_color="4F46E5", end_color="4F46E5", fill_type="solid")
            cell.alignment = Alignment(horizontal="center", vertical="center")
        
        # Write data
        for row_idx, row in enumerate(data, 2):
            for col_idx, key in enumerate(headers, 1):
                value = row[key]
                cell = ws.cell(row=row_idx, column=col_idx, value=value)
                
                # Color coding for sentiment labels
                if key == 'Duygu Etiketi':
                    color = _get_sentiment_color(row['Duygu Etiketi'])
                    cell.fill = PatternFill(start_color=color, end_color=color, fill_type="solid")
                    cell.font = Font(bold=True, color="FFFFFF")
                
                cell.alignment = Alignment(vertical="top", wrap_text=True)
        
        # Auto-adjust column widths
        for column in ws.columns:
            max_length = 0
            column_letter = column[0].column_letter
            for cell in column:
                try:
                    if len(str(cell.value)) > max_length:
                        max_length = len(str(cell.value))
                except:
                    pass
            adjusted_width = min(max_length + 2, 50)
            ws.column_dimensions[column_letter].width = adjusted_width
        
        # Add statistics sheet if requested
        if include_stats:
            _add_stats_sheet(wb, results, model_type)
        
        wb.save(file_path)
        return True
        
    except Exception as e:
        print(f"Excel export error: {e}")
        return False


def export_sentiment_to_word(
    file_path: str,
    results: List[Dict],
    model_type: str = "BERT",
    hybrid_results: Optional[List[Dict]] = None
) -> bool:
    """
    Export sentiment analysis results to Word (.docx) format.
    
    Args:
        file_path: Target file path
        results: List of sentiment result dictionaries
        model_type: "BERT" or "Online AI"
        hybrid_results: Optional hybrid comparison results
    """
    try:
        from docx import Document
        from docx.shared import Inches, Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_TABLE_ALIGNMENT
        
        doc = Document()
        
        # Title
        title = doc.add_heading('Duygu Analizi Raporu', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Metadata
        meta = doc.add_paragraph()
        meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
        meta.add_run(f'\nModel: {model_type}\n').italic = True
        meta.add_run(f'Tarih: {datetime.now().strftime("%d.%m.%Y %H:%M")}\n').italic = True
        meta.add_run(f'Toplam Belge: {len(results)}\n').italic = True
        
        doc.add_paragraph()
        
        # Summary statistics
        doc.add_heading('1. Özet İstatistikler', level=1)
        _add_summary_stats_to_doc(doc, results)
        
        # Detailed results
        doc.add_heading('2. Belge Bazlı Sonuçlar', level=1)
        
        for i, r in enumerate(results, 1):
            # Document heading
            p = doc.add_paragraph()
            p.add_run(f'{i}. {r.get("title", "Bilinmeyen")}').bold = True
            
            # Sentiment info
            label = _translate_label(r.get('label', 'neutral'))
            score = r.get('score', 0.5)
            summary = r.get('summary', '')
            
            info = doc.add_paragraph()
            info.add_run(f'Duygu: {label} | Güven: {score:.1%}')
            if summary:
                info.add_run(f' | {summary}')
            
            doc.add_paragraph()  # Spacing
        
        # Hybrid comparison if available
        if hybrid_results:
            doc.add_heading('3. Model Karşılaştırması', level=1)
            _add_hybrid_comparison_to_doc(doc, hybrid_results)
        
        # Footer
        doc.add_paragraph()
        footer = doc.add_paragraph()
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer.add_run('LexiScholar Akademik Analiz Yazılımı').italic = True
        
        doc.save(file_path)
        return True
        
    except Exception as e:
        print(f"Word export error: {e}")
        return False


def export_hybrid_sentiment_to_word(file_path: str, results: List[Dict], model_type: str = "AI") -> bool:
    """
    Export hybrid sentiment comparison to Word with both local and online model results.
    
    Args:
        file_path: Target file path
        results: List of hybrid results with 'local' and 'online' keys
        model_type: Name of the online AI model
    """
    try:
        from docx import Document
        from docx.shared import Inches, Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        
        doc = Document()
        
        # Title
        title = doc.add_heading('Hibrit Duygu Analizi Raporu', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Metadata
        meta = doc.add_paragraph()
        meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
        meta.add_run(f'Model: BERT + {model_type} | ').italic = True
        meta.add_run(f'Tarih: {datetime.now().strftime("%d.%m.%Y %H:%M")} | ').italic = True
        meta.add_run(f'Toplam: {len(results)} belge').italic = True
        
        doc.add_paragraph()
        
        # Calculate statistics
        match_count = 0
        for r in results:
            local = r.get('local', {})
            online = r.get('online', {})
            if local.get('label') == online.get('label'):
                match_count += 1
        
        match_rate = f"{(match_count/len(results)*100):.1f}%" if results else "0%"
        
        # Summary statistics
        stats_heading = doc.add_heading('1. Model Uyum İstatistikleri', level=1)
        
        stats = doc.add_paragraph()
        stats.add_run(f'Eşleşme Oranı: ').bold = True
        stats.add_run(f'{match_rate} ({match_count}/{len(results)} belge)')
        
        doc.add_paragraph()
        
        # Detailed comparison table
        doc.add_heading('2. Belge Bazlı Karşılaştırma', level=1)
        
        for i, r in enumerate(results, 1):
            local = r.get('local', {})
            online = r.get('online', {})
            
            doc_title = r.get('title', f'Belge {i}')
            
            # Document heading
            doc.add_heading(f'{i}. {doc_title}', level=2)
            
            # Create comparison table for this document
            table = doc.add_table(rows=3, cols=3)
            table.style = 'Table Grid'
            
            # Header row
            header_cells = table.rows[0].cells
            header_cells[0].text = 'Model'
            header_cells[1].text = 'Duygu'
            header_cells[2].text = 'Skor / Güven'
            
            # Make header bold
            for cell in header_cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True
            
            # BERT row
            bert_cells = table.rows[1].cells
            bert_cells[0].text = '🤖 BERT (Lokal)'
            l_label = _translate_label(local.get('label', 'neutral'))
            l_score = local.get('score', 0.5)
            bert_cells[1].text = l_label
            bert_cells[2].text = f'{l_score:.1%}'
            
            # AI row
            ai_cells = table.rows[2].cells
            ai_cells[0].text = f'🤖 {model_type} (Online)'
            o_label = _translate_label(online.get('label', 'neutral'))
            o_score = online.get('score', 0.5)
            o_conf = online.get('confidence', 0.5)
            ai_cells[1].text = o_label
            ai_cells[2].text = f'Skor: {o_score:.1%} | Güven: {o_conf:.1%}'
            
            # Color code the cells based on match
            is_match = local.get('label') == online.get('label')
            
            # Add summary
            doc.add_paragraph()
            summary_para = doc.add_paragraph()
            summary_para.add_run('BERT Özeti: ').bold = True
            summary_para.add_run(local.get('summary', 'Özet yok'))
            
            summary_para2 = doc.add_paragraph()
            summary_para2.add_run(f'{model_type} Özeti: ').bold = True
            summary_para2.add_run(online.get('summary', 'Özet yok'))
            
            # Match indicator
            match_para = doc.add_paragraph()
            if is_match:
                match_run = match_para.add_run('✓ Modeller uyumlu')
                match_run.font.color.rgb = RGBColor(0x10, 0xB9, 0x81)
                match_run.bold = True
            else:
                match_run = match_para.add_run('⚠ Modeller farklı sonuç verdi')
                match_run.font.color.rgb = RGBColor(0xF5, 0x9E, 0x0B)
                match_run.bold = True
            
            doc.add_paragraph()
        
        # Footer
        doc.add_paragraph()
        footer = doc.add_paragraph()
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer.add_run('LexiScholar Akademik Analiz Yazılımı').italic = True
        
        doc.save(file_path)
        return True
        
    except Exception as e:
        print(f"Hybrid Word export error: {e}")
        return False


def export_hybrid_sentiment_to_excel(file_path: str, results: List[Dict]) -> bool:
    """
    Export hybrid sentiment comparison to Excel.
    
    Results format:
    {
        "title": "Belge Adı",
        "local": {"label": "positive", "score": 0.8, "summary": "..."},
        "online": {"label": "very positive", "score": 0.9, "confidence": 0.95, "summary": "..."}
    }
    """
    try:
        import pandas as pd
        from openpyxl import Workbook
        from openpyxl.styles import Font, PatternFill, Alignment
        
        data = []
        for r in results:
            local = r.get('local', {})
            online = r.get('online', {})
            
            data.append({
                'Belge': r.get('title', 'Bilinmeyen'),
                'BERT Duygu': _translate_label(local.get('label', 'neutral')),
                'BERT Skor': local.get('score', 0.5),
                'AI Duygu': _translate_label(online.get('label', 'neutral')),
                'AI Skor': online.get('score', 0.5),
                'AI Güven': online.get('confidence', 0.5),
                'Uyum': 'Evet' if local.get('label') == online.get('label') else 'Hayır',
                'BERT Özet': local.get('summary', ''),
                'AI Özet': online.get('summary', '')
            })
        
        df = pd.DataFrame(data)
        
        # Create Excel with conditional formatting for matches
        wb = Workbook()
        ws = wb.active
        ws.title = "Model Karşılaştırması"
        
        # Write headers
        headers = list(df.columns)
        for col_idx, header in enumerate(headers, 1):
            cell = ws.cell(row=1, column=col_idx, value=header)
            cell.font = Font(bold=True, color="FFFFFF")
            cell.fill = PatternFill(start_color="4F46E5", end_color="4F46E5", fill_type="solid")
            cell.alignment = Alignment(horizontal="center", vertical="center")
        
        # Write data
        for row_idx, row in enumerate(data, 2):
            for col_idx, key in enumerate(headers, 1):
                value = row[key]
                cell = ws.cell(row=row_idx, column=col_idx, value=value)
                
                # Color code matches
                if key == 'Uyum':
                    if row[key] == 'Evet':
                        cell.fill = PatternFill(start_color="10B981", end_color="10B981", fill_type="solid")
                        cell.font = Font(color="FFFFFF", bold=True)
                    else:
                        cell.fill = PatternFill(start_color="EF4444", end_color="EF4444", fill_type="solid")
                        cell.font = Font(color="FFFFFF", bold=True)
                
                cell.alignment = Alignment(vertical="top", wrap_text=True)
        
        # Auto-adjust column widths
        for column in ws.columns:
            max_length = 0
            column_letter = column[0].column_letter
            for cell in column:
                try:
                    if len(str(cell.value)) > max_length:
                        max_length = len(str(cell.value))
                except:
                    pass
            adjusted_width = min(max_length + 2, 50)
            ws.column_dimensions[column_letter].width = adjusted_width
        
        wb.save(file_path)
        return True
        
    except Exception as e:
        print(f"Hybrid export error: {e}")
        return False


def export_sentiment_to_html(
    file_path: str,
    results: List[Dict],
    model_type: str = "BERT"
) -> bool:
    """
    Export sentiment analysis results to detailed HTML format.
    
    Args:
        file_path: Target file path
        results: List of sentiment result dictionaries
        model_type: "BERT" or "Online AI"
    """
    try:
        labels = [r.get('label', 'neutral') for r in results]
        scores = [r.get('score', 0.5) for r in results]
        stats = _compute_sentiment_stats(labels, scores)
        generated_at = datetime.now().strftime('%d.%m.%Y %H:%M')
        report_title = 'Duygu Analizi Raporu'
        meta_text = f"{generated_at} | Toplam: {len(results)} belge"

        rows_html = ""
        for i, r in enumerate(results, 1):
            label = r.get('label', 'neutral')
            score = r.get('score', 0.5)
            badge_color = _get_label_hex(label)
            progress_color = _score_color(score)
            rows_html += f"""
                <tr>
                    <td class="idx">#{i}</td>
                    <td class="doc">{r.get('title', 'Bilinmeyen')}</td>
                    <td><span class="badge" style="background:{badge_color}18;color:{badge_color};border-color:{badge_color}33">{_translate_label(label)}</span></td>
                    <td>
                        <div class="score">{score:.0%}</div>
                        <div class="bar"><div class="fill" style="width:{score:.0%};background:{progress_color}"></div></div>
                    </td>
                    <td class="summary">{r.get('summary', 'Analiz özeti mevcut değil.')}</td>
                </tr>
            """

        html_content = _build_html_shell(
            page_title=report_title,
            header_title=report_title,
            meta_text=meta_text,
            stat_cards=[
                ("Pozitif", str(stats['pos_total']), stats['pos_pct'], "positive"),
                ("Nötr", str(stats['neu_total']), stats['neu_pct'], "neutral"),
                ("Negatif", str(stats['neg_total']), stats['neg_pct'], "negative"),
                ("Ortalama Skor", stats['avg_percentage'], stats['overall_sentiment'], "accent"),
            ],
            section_title="Belge Bazlı Sonuçlar",
            section_body=f"""
            <table class="result-table single-mode">
                <thead>
                    <tr>
                        <th>Sıra</th>
                        <th>Belge</th>
                        <th>Duygu</th>
                        <th>Skor</th>
                        <th>Özet / Analiz</th>
                    </tr>
                </thead>
                <tbody>
                    {rows_html}
                </tbody>
            </table>
            """,
            footer_text=f"LexiScholar Akademik Analiz Yazılımı | {datetime.now().strftime('%d.%m.%Y')}"
        )

        with open(file_path, 'w', encoding='utf-8') as f:
            f.write(html_content)

        return True

    except Exception as e:
        print(f"HTML export error: {e}")
        return False


def export_hybrid_sentiment_to_html(file_path: str, results: List[Dict], model_type: str = "AI") -> bool:
    """
    Export hybrid sentiment comparison to detailed HTML format.
    Shows both BERT and AI model results with summaries.
    
    Args:
        file_path: Target file path
        results: List of hybrid results with 'local' and 'online' keys
        model_type: Name of the online AI model
    """
    try:
        generated_at = datetime.now().strftime('%d.%m.%Y %H:%M')
        local_scores = [r.get('local', {}).get('score', 0.5) for r in results]
        online_scores = [r.get('online', {}).get('score', 0.5) for r in results]

        exact_count = 0
        close_count = 0
        different_count = 0
        rows_html = ""

        for i, r in enumerate(results, 1):
            local = r.get('local', {})
            online = r.get('online', {})
            comparison = _compare_sentiment_labels(local.get('label', 'neutral'), online.get('label', 'neutral'))

            if comparison['state'] == 'exact':
                exact_count += 1
            elif comparison['state'] == 'close':
                close_count += 1
            else:
                different_count += 1

            l_label = local.get('label', 'neutral')
            o_label = online.get('label', 'neutral')
            l_score = local.get('score', 0.5)
            o_score = online.get('score', 0.5)
            o_conf = online.get('confidence', 0.5)
            l_color = _get_label_hex(l_label)
            o_color = _get_label_hex(o_label)
            compare_color = _comparison_color(comparison['state'])

            rows_html += f"""
                <tr>
                    <td class="idx">#{i}</td>
                    <td class="doc">{r.get('title', 'Bilinmeyen')}</td>
                    <td class="model-cell">
                        <span class="badge" style="background:{l_color}18;color:{l_color};border-color:{l_color}33">{_translate_label(l_label)}</span>
                        <div class="score">{l_score:.0%}</div>
                        <div class="bar"><div class="fill" style="width:{l_score:.0%};background:{_score_color(l_score)}"></div></div>
                        <div class="summary">{local.get('summary', 'Özet yok')}</div>
                    </td>
                    <td class="model-cell">
                        <span class="badge" style="background:{o_color}18;color:{o_color};border-color:{o_color}33">{_translate_label(o_label)}</span>
                        <div class="score">{o_score:.0%}</div>
                        <div class="bar"><div class="fill" style="width:{o_score:.0%};background:{_score_color(o_score)}"></div></div>
                        <div class="meta-inline">Güven: {o_conf:.0%}</div>
                        <div class="summary">{online.get('summary', 'Özet yok')}</div>
                    </td>
                    <td><span class="badge compare" style="background:{compare_color}18;color:{compare_color};border-color:{compare_color}33">{comparison['label']}</span></td>
                </tr>
            """

        avg_local = sum(local_scores) / len(local_scores) if local_scores else 0.5
        avg_online = sum(online_scores) / len(online_scores) if online_scores else 0.5
        html_content = _build_html_shell(
            page_title='Hibrit Duygu Analizi Raporu',
            header_title='Hibrit Duygu Analizi Raporu',
            meta_text=f"{generated_at} | Toplam: {len(results)} belge",
            stat_cards=[
                ("Uyumlu", str(exact_count), _ratio_text(exact_count, len(results)), "positive"),
                ("Yakın", str(close_count), _ratio_text(close_count, len(results)), "warning"),
                ("Farklı", str(different_count), _ratio_text(different_count, len(results)), "negative"),
                ("Ort. Skorlar", f"{avg_local:.0%} / {avg_online:.0%}", "BERT / Online", "accent"),
            ],
            section_title='Belge Bazlı Karşılaştırma',
            section_body=f"""
            <table class="result-table hybrid-mode">
                <thead>
                    <tr>
                        <th>Sıra</th>
                        <th>Belge</th>
                        <th>Lokal</th>
                        <th>Online</th>
                        <th>Karşılaştırma</th>
                    </tr>
                </thead>
                <tbody>
                    {rows_html}
                </tbody>
            </table>
            """,
            footer_text=f"LexiScholar Akademik Analiz Yazılımı | {datetime.now().strftime('%d.%m.%Y')}"
        )

        with open(file_path, 'w', encoding='utf-8') as f:
            f.write(html_content)

        return True

    except Exception as e:
        print(f"Hybrid HTML export error: {e}")
        return False


# Helper functions
def _build_html_shell(page_title: str, header_title: str, meta_text: str, stat_cards: List[tuple], section_title: str, section_body: str, footer_text: str) -> str:
    cards_html = "".join(
        f"""
        <div class="stat-card {variant}">
            <div class="stat-value">{value}</div>
            <div class="stat-title">{title}</div>
            <div class="stat-subtitle">{subtitle}</div>
        </div>
        """
        for title, value, subtitle, variant in stat_cards
    )
    return f"""<!DOCTYPE html>
<html lang="tr">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>{page_title}</title>
    <style>
        :root {{
            --bg: #eef2f7;
            --panel: #ffffff;
            --panel-soft: #f8fafc;
            --text: #0f172a;
            --muted: #64748b;
            --line: #dbe3ee;
            --brand: #4f46e5;
            --brand-2: #7c3aed;
            --positive: #059669;
            --warning: #d97706;
            --negative: #dc2626;
            --neutral: #475569;
        }}
        * {{ box-sizing: border-box; }}
        body {{
            margin: 0;
            background: var(--bg);
            color: var(--text);
            font-family: 'Segoe UI', Tahoma, Geneva, Verdana, sans-serif;
            padding: 20px;
        }}
        .container {{
            max-width: 1180px;
            margin: 0 auto;
            background: var(--panel);
            border: 1px solid var(--line);
            border-radius: 18px;
            overflow: hidden;
            box-shadow: 0 12px 36px rgba(15, 23, 42, 0.08);
        }}
        .header {{
            background: linear-gradient(135deg, var(--brand), var(--brand-2));
            color: #fff;
            padding: 24px 28px 20px;
        }}
        .header h1 {{
            margin: 0 0 6px 0;
            font-size: 30px;
            line-height: 1.2;
        }}
        .meta {{
            font-size: 14px;
            color: rgba(255,255,255,0.92);
        }}
        .stats {{
            display: grid;
            grid-template-columns: repeat(auto-fit, minmax(150px, 1fr));
            gap: 12px;
            padding: 18px 20px;
            background: var(--panel-soft);
            border-bottom: 1px solid var(--line);
        }}
        .stat-card {{
            border-radius: 14px;
            padding: 16px 14px;
            color: #fff;
            min-height: 96px;
        }}
        .stat-card.positive {{ background: linear-gradient(135deg, #10b981, #059669); }}
        .stat-card.warning {{ background: linear-gradient(135deg, #f59e0b, #d97706); }}
        .stat-card.negative {{ background: linear-gradient(135deg, #ef4444, #dc2626); }}
        .stat-card.neutral {{ background: linear-gradient(135deg, #64748b, #475569); }}
        .stat-card.accent {{ background: linear-gradient(135deg, #3b82f6, #4f46e5); }}
        .stat-value {{ font-size: 30px; font-weight: 800; line-height: 1.1; }}
        .stat-title {{ font-size: 12px; text-transform: uppercase; letter-spacing: 0.08em; margin-top: 8px; opacity: 0.95; }}
        .stat-subtitle {{ font-size: 12px; margin-top: 4px; opacity: 0.9; }}
        .section {{ padding: 22px 20px 20px; }}
        .section h2 {{ margin: 0 0 14px 0; font-size: 24px; }}
        .result-table {{ width: 100%; border-collapse: separate; border-spacing: 0; table-layout: fixed; }}
        .result-table th {{
            text-align: left;
            font-size: 12px;
            text-transform: uppercase;
            letter-spacing: 0.08em;
            color: var(--muted);
            background: var(--panel-soft);
            border-bottom: 1px solid var(--line);
            padding: 12px 14px;
        }}
        .result-table td {{
            padding: 14px;
            border-bottom: 1px solid var(--line);
            vertical-align: top;
            background: #fff;
        }}
        .result-table tr:nth-child(even) td {{ background: #fcfdff; }}
        .idx {{ width: 72px; color: var(--muted); font-weight: 700; }}
        .doc {{ font-weight: 700; }}
        .badge {{
            display: inline-flex;
            align-items: center;
            gap: 6px;
            padding: 6px 10px;
            border-radius: 999px;
            border: 1px solid;
            font-size: 12px;
            font-weight: 700;
            white-space: nowrap;
        }}
        .badge.compare {{ min-width: 78px; justify-content: center; }}
        .score {{ font-size: 18px; font-weight: 800; margin: 8px 0 6px; }}
        .bar {{ width: 100%; height: 8px; background: #e2e8f0; border-radius: 999px; overflow: hidden; }}
        .fill {{ height: 100%; border-radius: 999px; }}
        .summary {{ margin-top: 8px; color: var(--muted); line-height: 1.5; font-size: 13px; }}
        .meta-inline {{ margin-top: 8px; color: var(--muted); font-size: 12px; font-weight: 600; }}
        .model-cell {{ min-width: 0; }}
        .footer {{ padding: 16px 20px; background: #111827; color: #e5e7eb; font-size: 12px; text-align: center; }}
        @media (max-width: 900px) {{
            .result-table, .result-table thead, .result-table tbody, .result-table th, .result-table td, .result-table tr {{ display: block; }}
            .result-table thead {{ display: none; }}
            .result-table td {{ border-bottom: none; padding: 10px 14px; }}
            .result-table tr {{ border-bottom: 1px solid var(--line); padding: 6px 0; }}
        }}
    </style>
</head>
<body>
    <div class="container">
        <div class="header">
            <h1>{header_title}</h1>
            <div class="meta">{meta_text}</div>
        </div>
        <div class="stats">{cards_html}</div>
        <div class="section">
            <h2>{section_title}</h2>
            {section_body}
        </div>
        <div class="footer">{footer_text}</div>
    </div>
</body>
</html>"""


def _compute_sentiment_stats(labels: List[str], scores: List[float]) -> Dict:
    label_counts = {}
    for label in labels:
        label_counts[label] = label_counts.get(label, 0) + 1
    total = len(labels)
    pos_total = label_counts.get("very positive", 0) + label_counts.get("positive", 0)
    neg_total = label_counts.get("very negative", 0) + label_counts.get("negative", 0)
    neu_total = label_counts.get("neutral", 0)
    avg_score = sum(scores) / len(scores) if scores else 0.5
    if pos_total > neg_total and pos_total > neu_total:
        overall_sentiment = "Genel olumlu"
    elif neg_total > pos_total and neg_total > neu_total:
        overall_sentiment = "Genel olumsuz"
    else:
        overall_sentiment = "Genel nötr"
    return {
        'pos_total': pos_total,
        'neg_total': neg_total,
        'neu_total': neu_total,
        'pos_pct': _ratio_text(pos_total, total),
        'neg_pct': _ratio_text(neg_total, total),
        'neu_pct': _ratio_text(neu_total, total),
        'avg_percentage': f"{avg_score:.0%}",
        'overall_sentiment': overall_sentiment,
    }


def _ratio_text(part: int, total: int) -> str:
    return f"{(part / total * 100):.1f}%" if total else "0.0%"


def _get_label_hex(label: str) -> str:
    mapping = {
        'very positive': '#059669',
        'positive': '#10B981',
        'neutral': '#64748B',
        'negative': '#EF4444',
        'very negative': '#B91C1C',
        'mixed': '#D97706',
        'error': '#94A3B8'
    }
    return mapping.get(label, '#64748B')


def _score_color(score: float) -> str:
    if score >= 0.7:
        return '#10B981'
    if score >= 0.4:
        return '#64748B'
    return '#EF4444'


def _compare_sentiment_labels(label_a: str, label_b: str) -> Dict:
    if label_a == label_b:
        return {'state': 'exact', 'label': 'Uyumlu'}
    sentiment_scale = {
        'very negative': 0,
        'negative': 1,
        'neutral': 2,
        'positive': 3,
        'very positive': 4,
    }
    if label_a in sentiment_scale and label_b in sentiment_scale:
        if abs(sentiment_scale[label_a] - sentiment_scale[label_b]) == 1:
            return {'state': 'close', 'label': 'Yakın'}
    return {'state': 'different', 'label': 'Farklı'}


def _comparison_color(state: str) -> str:
    if state == 'exact':
        return '#059669'
    if state == 'close':
        return '#D97706'
    return '#DC2626'


def _translate_label(label: str) -> str:
    """Translate sentiment labels to Turkish."""
    translations = {
        'very positive': 'Çok Pozitif',
        'positive': 'Pozitif',
        'neutral': 'Nötr',
        'negative': 'Negatif',
        'very negative': 'Çok Negatif',
        'mixed': 'Karışık',
        'error': 'Hata'
    }
    return translations.get(label, label.capitalize())


def _get_sentiment_color(label: str) -> str:
    """Get hex color for sentiment label."""
    colors = {
        'Çok Pozitif': '059669',
        'Pozitif': '10B981',
        'Nötr': '64748B',
        'Negatif': 'EF4444',
        'Çok Negatif': 'B91C1C',
        'Karışık': 'F59E0B',
        'Hata': '94A3B8'
    }
    return colors.get(label, '64748B')


def _add_stats_sheet(wb, results, model_type):
    """Add statistics summary sheet to Excel workbook."""
    from openpyxl.styles import Font, PatternFill
    
    ws = wb.create_sheet(title="İstatistikler")
    
    # Calculate stats
    labels = [r.get('label', 'neutral') for r in results]
    label_counts = {}
    for l in labels:
        label_counts[l] = label_counts.get(l, 0) + 1
    
    # Write stats
    ws.cell(row=1, column=1, value="İstatistik").font = Font(bold=True)
    ws.cell(row=1, column=2, value="Değer").font = Font(bold=True)
    
    row = 2
    ws.cell(row=row, column=1, value="Toplam Belge")
    ws.cell(row=row, column=2, value=len(results))
    row += 1
    
    ws.cell(row=row, column=1, value="Model")
    ws.cell(row=row, column=2, value=model_type)
    row += 2
    
    ws.cell(row=row, column=1, value="Duygu Dağılımı").font = Font(bold=True)
    row += 1
    
    for label, count in sorted(label_counts.items()):
        ws.cell(row=row, column=1, value=_translate_label(label))
        ws.cell(row=row, column=2, value=count)
        row += 1


def _add_summary_stats_to_doc(doc, results):
    """Add summary statistics to Word document."""
    from docx.shared import Inches
    
    labels = [r.get('label', 'neutral') for r in results]
    label_counts = {}
    for l in labels:
        label_counts[l] = label_counts.get(l, 0) + 1
    
    # Create summary table
    table = doc.add_table(rows=len(label_counts) + 1, cols=2)
    table.style = 'Light Grid Accent 1'
    
    # Header
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Duygu'
    hdr_cells[1].text = 'Sayı'
    
    # Data
    for i, (label, count) in enumerate(sorted(label_counts.items()), 1):
        row_cells = table.rows[i].cells
        row_cells[0].text = _translate_label(label)
        row_cells[1].text = str(count)


def _add_hybrid_comparison_to_doc(doc, hybrid_results):
    """Add hybrid model comparison to Word document."""
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Light Grid Accent 1'
    
    # Header
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Belge'
    hdr_cells[1].text = 'BERT'
    hdr_cells[2].text = 'Online AI'
    hdr_cells[3].text = 'Uyum'
    
    # Data
    for r in hybrid_results:
        local = r.get('local', {})
        online = r.get('online', {})
        is_match = local.get('label') == online.get('label')
        
        row_cells = table.add_row().cells
        row_cells[0].text = r.get('title', 'Bilinmeyen')
        row_cells[1].text = _translate_label(local.get('label', 'neutral'))
        row_cells[2].text = _translate_label(online.get('label', 'neutral'))
        row_cells[3].text = '✓' if is_match else '✗'


def get_sentiment_export_formats() -> List[Dict]:
    """Return available export formats for sentiment analysis."""
    formats = [
        {
            'name': 'Excel Belgesi',
            'extension': 'xlsx',
            'filter': 'Excel Dosyaları (*.xlsx)',
            'description': 'Yapılandırılmış tablo formatında export'
        },
        {
            'name': 'Word Belgesi',
            'extension': 'docx',
            'filter': 'Word Belgeleri (*.docx)',
            'description': 'Akademik rapor formatında export'
        },
        {
            'name': 'HTML Rapor',
            'extension': 'html',
            'filter': 'HTML Dosyaları (*.html)',
            'description': 'Web tarayıcısında görüntülenebilir rapor'
        },
        {
            'name': 'JSON Verisi',
            'extension': 'json',
            'filter': 'JSON Dosyaları (*.json)',
            'description': 'Programatik analiz için ham veri'
        }
    ]
    return formats
