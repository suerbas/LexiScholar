"""
Topic Modeling Export Module
Provides Excel, Word, and HTML export for topic modeling results.
"""

import os
from datetime import datetime
from typing import List, Dict, Optional
from pathlib import Path

def export_topics_to_excel(
    file_path: str,
    topic_data: Dict,
    model_type: str = "LDA"
) -> bool:
    """
    Export topic modeling results to Excel (.xlsx) format.
    
    Args:
        file_path: Target file path
        topic_data: Topic modeling result dictionary
        model_type: "LDA" or model name
    """
    try:
        import pandas as pd
        from openpyxl import Workbook
        from openpyxl.styles import Font, PatternFill, Alignment
        
        wb = Workbook()
        
        # Topics sheet
        ws_topics = wb.active
        ws_topics.title = "Konular"
        
        headers = ["Konu ID", "Etiket", "Anahtar Kelimeler"]
        for col_idx, header in enumerate(headers, 1):
            cell = ws_topics.cell(row=1, column=col_idx, value=header)
            cell.font = Font(bold=True, color="FFFFFF")
            cell.fill = PatternFill(start_color="4F46E5", end_color="4F46E5", fill_type="solid")
            cell.alignment = Alignment(horizontal="center", vertical="center")
        
        topics = topic_data.get("topics", [])
        for row_idx, topic in enumerate(topics, 2):
            ws_topics.cell(row=row_idx, column=1, value=topic.get("id", row_idx-2))
            ws_topics.cell(row=row_idx, column=2, value=topic.get("label", f"Konu {row_idx-1}"))
            words = topic.get("words", [])
            words_str = ", ".join([f"{w} ({s:.3f})" for w, s in words[:10]])
            ws_topics.cell(row=row_idx, column=3, value=words_str)
            ws_topics.cell(row=row_idx, column=3).alignment = Alignment(vertical="top", wrap_text=True)
        
        # Document topics sheet
        ws_docs = wb.create_sheet(title="Belge Konuları")
        doc_headers = ["Belge", "Baskın Konu", "Konu Ağırlıkları"]
        for col_idx, header in enumerate(doc_headers, 1):
            cell = ws_docs.cell(row=1, column=col_idx, value=header)
            cell.font = Font(bold=True, color="FFFFFF")
            cell.fill = PatternFill(start_color="4F46E5", end_color="4F46E5", fill_type="solid")
            cell.alignment = Alignment(horizontal="center", vertical="center")
        
        doc_topics = topic_data.get("doc_topics", [])
        for row_idx, doc in enumerate(doc_topics, 2):
            ws_docs.cell(row=row_idx, column=1, value=doc.get("title", "Bilinmeyen"))
            dominant = doc.get("dominant_topic", 0)
            ws_docs.cell(row=row_idx, column=2, value=f"Konu {dominant + 1}")
            weights = doc.get("topic_weights", [])
            weights_str = ", ".join([f"{w:.2f}" for w in weights])
            ws_docs.cell(row=row_idx, column=3, value=weights_str)
        
        # Auto-adjust column widths
        for ws in [ws_topics, ws_docs]:
            for column in ws.columns:
                max_length = 0
                column_letter = column[0].column_letter
                for cell in column:
                    try:
                        if len(str(cell.value)) > max_length:
                            max_length = len(str(cell.value))
                    except:
                        pass
                adjusted_width = min(max_length + 2, 60)
                ws.column_dimensions[column_letter].width = adjusted_width
        
        wb.save(file_path)
        return True
        
    except Exception as e:
        print(f"Excel export error: {e}")
        return False


def export_topics_to_word(
    file_path: str,
    topic_data: Dict,
    model_type: str = "LDA"
) -> bool:
    """
    Export topic modeling results to Word (.docx) format.
    
    Args:
        file_path: Target file path
        topic_data: Topic modeling result dictionary
        model_type: "LDA" or model name
    """
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        
        doc = Document()
        
        # Title
        title = doc.add_heading('Konu Modelleme Raporu', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Metadata
        meta = doc.add_paragraph()
        meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
        meta.add_run(f'Model: {model_type}\n').italic = True
        meta.add_run(f'Tarih: {datetime.now().strftime("%d.%m.%Y %H:%M")}\n').italic = True
        
        topics = topic_data.get("topics", [])
        doc_topics = topic_data.get("doc_topics", [])
        
        meta.add_run(f'Toplam Konu: {len(topics)}\n').italic = True
        meta.add_run(f'Toplam Belge: {len(doc_topics)}\n').italic = True
        
        doc.add_paragraph()
        
        # Topics section
        doc.add_heading('1. Keşfedilen Konular', level=1)
        
        for i, topic in enumerate(topics, 1):
            p = doc.add_paragraph()
            p.add_run(f'{i}. {topic.get("label", f"Konu {i}")}').bold = True
            
            words = topic.get("words", [])
            words_str = ", ".join([f"{w} ({s:.3f})" for w, s in words[:10]])
            doc.add_paragraph(f'Anahtar kelimeler: {words_str}')
        
        # Document distribution section
        doc.add_heading('2. Belge-Konu Dağılımı', level=1)
        
        table = doc.add_table(rows=1, cols=3)
        table.style = 'Light Grid Accent 1'
        
        # Header
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Belge'
        hdr_cells[1].text = 'Baskın Konu'
        hdr_cells[2].text = 'Konu Ağırlıkları'
        
        # Data
        for doc_t in doc_topics:
            row_cells = table.add_row().cells
            row_cells[0].text = doc_t.get("title", "Bilinmeyen")
            dominant = doc_t.get("dominant_topic", 0)
            row_cells[1].text = f"Konu {dominant + 1}"
            weights = doc_t.get("topic_weights", [])
            row_cells[2].text = ", ".join([f"{w:.2f}" for w in weights])
        
        # Footer
        doc.add_paragraph()
        footer = doc.add_paragraph()
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer.add_run('LexiScholar Akademik Analiz Yazılımı').italic = True
        
        doc.save(file_path)
        return True
        
    except Exception as e:
        print(f"Word export error: {e}")
        return False


def export_hybrid_topics_to_excel(file_path: str, topic_data: Dict) -> bool:
    """
    Export hybrid topic modeling comparison to Excel.
    """
    try:
        import pandas as pd
        from openpyxl import Workbook
        from openpyxl.styles import Font, PatternFill, Alignment
        
        wb = Workbook()
        
        local = topic_data.get("local", {})
        online = topic_data.get("online", {})
        comparison = topic_data.get("comparison", {})
        
        # Local topics sheet
        ws_local = wb.active
        ws_local.title = "LDA Konular"
        
        headers = ["Konu ID", "Etiket", "Anahtar Kelimeler"]
        for col_idx, header in enumerate(headers, 1):
            cell = ws_local.cell(row=1, column=col_idx, value=header)
            cell.font = Font(bold=True, color="FFFFFF")
            cell.fill = PatternFill(start_color="0369a1", end_color="0369a1", fill_type="solid")
        
        local_topics = local.get("topics", [])
        for row_idx, topic in enumerate(local_topics, 2):
            ws_local.cell(row=row_idx, column=1, value=topic.get("id", row_idx-2))
            ws_local.cell(row=row_idx, column=2, value=topic.get("label", f"Konu {row_idx-1}"))
            words = topic.get("words", [])
            words_str = ", ".join([f"{w} ({s:.3f})" for w, s in words[:10]])
            ws_local.cell(row=row_idx, column=3, value=words_str)
        
        # Online topics sheet
        ws_online = wb.create_sheet(title="AI Konular")
        for col_idx, header in enumerate(headers, 1):
            cell = ws_online.cell(row=1, column=col_idx, value=header)
            cell.font = Font(bold=True, color="FFFFFF")
            cell.fill = PatternFill(start_color="7c3aed", end_color="7c3aed", fill_type="solid")
        
        online_topics = online.get("topics", [])
        for row_idx, topic in enumerate(online_topics, 2):
            ws_online.cell(row=row_idx, column=1, value=topic.get("id", row_idx-2))
            ws_online.cell(row=row_idx, column=2, value=topic.get("label", f"Konu {row_idx-1}"))
            words = topic.get("words", [])
            words_str = ", ".join([f"{w} ({s:.3f})" for w, s in words[:10]])
            ws_online.cell(row=row_idx, column=3, value=words_str)
        
        # Comparison sheet
        ws_comp = wb.create_sheet(title="Karşılaştırma")
        comp_headers = ["Belge", "LDA Baskın", "AI Baskın", "Durum"]
        for col_idx, header in enumerate(comp_headers, 1):
            cell = ws_comp.cell(row=1, column=col_idx, value=header)
            cell.font = Font(bold=True, color="FFFFFF")
            cell.fill = PatternFill(start_color="4F46E5", end_color="4F46E5", fill_type="solid")
        
        doc_differences = comparison.get("doc_differences", [])
        for row_idx, diff in enumerate(doc_differences, 2):
            ws_comp.cell(row=row_idx, column=1, value=diff.get("title", "Bilinmeyen"))
            ws_comp.cell(row=row_idx, column=2, value=f"Konu {diff.get('local_dominant', 0) + 1}")
            ws_comp.cell(row=row_idx, column=3, value=f"Konu {diff.get('online_dominant', 0) + 1}")
            status = diff.get("status", "farklı")
            status_cell = ws_comp.cell(row=row_idx, column=4, value=status.capitalize())
            
            if status == "uyumlu":
                status_cell.fill = PatternFill(start_color="10B981", end_color="10B981", fill_type="solid")
                status_cell.font = Font(color="FFFFFF", bold=True)
            elif status == "yakın":
                status_cell.fill = PatternFill(start_color="F59E0B", end_color="F59E0B", fill_type="solid")
                status_cell.font = Font(color="FFFFFF", bold=True)
            else:
                status_cell.fill = PatternFill(start_color="EF4444", end_color="EF4444", fill_type="solid")
                status_cell.font = Font(color="FFFFFF", bold=True)
        
        # Summary sheet
        ws_summary = wb.create_sheet(title="Özet")
        summary = comparison.get("summary", {})
        ws_summary.cell(row=1, column=1, value="Durum")
        ws_summary.cell(row=1, column=2, value="Sayı")
        ws_summary.cell(row=1, column=1).font = Font(bold=True)
        ws_summary.cell(row=1, column=2).font = Font(bold=True)
        
        row = 2
        for status, count in summary.items():
            ws_summary.cell(row=row, column=1, value=status.capitalize())
            ws_summary.cell(row=row, column=2, value=count)
            row += 1
        
        wb.save(file_path)
        return True
        
    except Exception as e:
        print(f"Hybrid Excel export error: {e}")
        return False


def export_hybrid_topics_to_word(file_path: str, topic_data: Dict, model_type: str = "AI") -> bool:
    """
    Export hybrid topic modeling comparison to Word.
    """
    try:
        from docx import Document
        from docx.shared import RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        
        doc = Document()
        
        # Title
        title = doc.add_heading('Hibrit Konu Modelleme Raporu', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Metadata
        local = topic_data.get("local", {})
        online = topic_data.get("online", {})
        comparison = topic_data.get("comparison", {})
        
        meta = doc.add_paragraph()
        meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
        meta.add_run(f'Model: LDA + {model_type}\n').italic = True
        meta.add_run(f'Tarih: {datetime.now().strftime("%d.%m.%Y %H:%M")}\n').italic = True
        
        # Summary
        summary = comparison.get("summary", {})
        meta.add_run(f'\nUyumlu: {summary.get("uyumlu", 0)} | Yakın: {summary.get("yakın", 0)} | Farklı: {summary.get("farklı", 0)}\n').italic = True
        
        doc.add_paragraph()
        
        # LDA Topics
        doc.add_heading('1. LDA (Lokal) Konular', level=1)
        local_topics = local.get("topics", [])
        for i, topic in enumerate(local_topics, 1):
            p = doc.add_paragraph()
            p.add_run(f'{i}. {topic.get("label", f"Konu {i}")}').bold = True
            words = topic.get("words", [])
            words_str = ", ".join([f"{w} ({s:.3f})" for w, s in words[:10]])
            doc.add_paragraph(f'Anahtar kelimeler: {words_str}')
        
        # AI Topics
        doc.add_heading(f'2. {model_type} (Online) Konular', level=1)
        online_topics = online.get("topics", [])
        for i, topic in enumerate(online_topics, 1):
            p = doc.add_paragraph()
            p.add_run(f'{i}. {topic.get("label", f"Konu {i}")}').bold = True
            words = topic.get("words", [])
            words_str = ", ".join([f"{w} ({s:.3f})" for w, s in words[:10]])
            doc.add_paragraph(f'Anahtar kelimeler: {words_str}')
        
        # Comparison table
        doc.add_heading('3. Belge Bazlı Karşılaştırma', level=1)
        
        table = doc.add_table(rows=1, cols=4)
        table.style = 'Light Grid Accent 1'
        
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Belge'
        hdr_cells[1].text = 'LDA'
        hdr_cells[2].text = 'AI'
        hdr_cells[3].text = 'Durum'
        
        doc_differences = comparison.get("doc_differences", [])
        for diff in doc_differences:
            row_cells = table.add_row().cells
            row_cells[0].text = diff.get("title", "Bilinmeyen")
            row_cells[1].text = f"Konu {diff.get('local_dominant', 0) + 1}"
            row_cells[2].text = f"Konu {diff.get('online_dominant', 0) + 1}"
            status = diff.get("status", "farklı")
            row_cells[3].text = status.capitalize()
        
        # Footer
        doc.add_paragraph()
        footer = doc.add_paragraph()
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer.add_run('LexiScholar Akademik Analiz Yazılımı').italic = True
        
        doc.save(file_path)
        return True
        
    except Exception as e:
        print(f"Hybrid Word export error: {e}")
        return False


def _translate_topic_label(label: str) -> str:
    """Translate topic labels if needed."""
    return label
