"""
DOCX Processor for LexiScholar
Extracts text from Word documents while preserving paragraph structure.
"""

from dataclasses import dataclass
from pathlib import Path
from typing import List, Optional
import os

try:
    from docx import Document
    from docx.opc.exceptions import PackageNotFoundError
except ImportError:
    Document = None
    PackageNotFoundError = Exception


@dataclass
class DocxExtractionResult:
    """Result of DOCX text extraction."""
    full_text: str
    paragraphs: List[str]
    success: bool
    error: Optional[str] = None


def extract_text(docx_path: str) -> DocxExtractionResult:
    """
    Extract text from a DOCX file.
    
    Preserves paragraph structure by separating paragraphs with double newlines.
    This maintains the logical structure needed for qualitative coding.
    
    Args:
        docx_path: Path to the DOCX file
        
    Returns:
        DocxExtractionResult with full text and individual paragraphs
    """
    if Document is None:
        return DocxExtractionResult(
            full_text="",
            success=False,
            error="python-docx library not installed"
        )
    
    # Memory Management: Check file size (limit 50MB)
    MAX_SIZE = 50 * 1024 * 1024
    if os.path.exists(docx_path) and os.path.getsize(docx_path) > MAX_SIZE:
        return DocxExtractionResult(
            full_text="",
            paragraphs=[],
            success=False,
            error="Dosya çok büyük (Limit: 50MB). Lütfen daha küçük parçalara bölün."
        )

    try:
        doc = Document(docx_path)
        paragraphs = []
        
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:  # Skip empty paragraphs
                paragraphs.append(text)
        
        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_texts = []
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text:
                        row_texts.append(cell_text)
                if row_texts:
                    paragraphs.append(" | ".join(row_texts))
        
        full_text = "\n\n".join(paragraphs)
        
        return DocxExtractionResult(
            full_text=full_text,
            paragraphs=paragraphs,
            success=True
        )
        
    except PackageNotFoundError:
        return DocxExtractionResult(
            full_text="",
            paragraphs=[],
            success=False,
            error="File not found or not a valid DOCX file"
        )
    except Exception as e:
        return DocxExtractionResult(
            full_text="",
            paragraphs=[],
            success=False,
            error=str(e)
        )


if __name__ == "__main__":
    import sys
    if len(sys.argv) > 1:
        result = extract_text(sys.argv[1])
        if result.success:
            print(f"Extracted {len(result.paragraphs)} paragraphs")
            print(f"Total characters: {len(result.full_text)}")
            print("\n--- First 500 characters ---")
            print(result.full_text[:500])
        else:
            print(f"Error: {result.error}")
    else:
        print("Usage: python docx_processor.py <docx_file>")
