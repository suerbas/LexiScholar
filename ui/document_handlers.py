"""
Document-related event handlers for LexiScholar.
"""

from PyQt6.QtWidgets import QMessageBox, QFileDialog, QDialog, QVBoxLayout, QLabel, QPushButton, QHBoxLayout, QWidget
from PyQt6.QtCore import Qt, QUrl
from pathlib import Path
import os
import tempfile
import atexit
import logging
from .common_ui import show_info, show_warning, show_error, ask_confirmation

logger = logging.getLogger(__name__)

class DocumentHandlersMixin:
    """Mixin for document management handlers."""
    
    def _on_document_selected(self, doc_id: int):
        """Handle document selection."""
        doc = self.doc_dao.get_by_id(doc_id)
        if doc:
            # Use current coder ID for segment and memo viewing
            coder_id = getattr(self, 'current_coder_id', 1)
            segments = self.segment_dao.get_by_document(doc_id, coder_id=coder_id)
            memos = self.memo_dao.get_by_document(doc_id, coder_id=coder_id)
            self.document_browser.set_document(
                doc_id, 
                doc['extracted_text'] or "", 
                segments, 
                memos,
                title=doc['title']
            )
            self.statusbar.showMessage(f"Belge: {doc['title']}")
            
            # Handle Audio Sync for Transcriptions
            if doc['file_type'] == "transcription" and doc.get('file_path'):
                self.audio_player.load_audio(doc['file_path'])
                self.audio_player.show()
            else:
                self.audio_player.hide()
                self.audio_player.player.stop()
            
            # Switch to document view if visualization is active
            # (Now handled by standalone windows, so no stack switching needed)
            pass
    
    def _on_document_deleted(self, doc_id: int):
        """Handle document deletion."""
        if hasattr(self, 'document_browser') and self.document_browser._current_doc_id == doc_id:
            self.document_browser.clear()
            self.statusbar.showMessage("Belge silindi.")
        self.set_dirty()
    
    def _on_chat_requested(self, doc_id: int, doc_title: str):
        """Open the AI chat in a central tab for the selected document."""
        from .chat_with_document_dialog import ChatWidget
        
        # Check if tab already exists
        tab_name = f"💬 AI: {doc_title}"
        for i in range(self.central_tabs.count()):
            if self.central_tabs.tabText(i) == tab_name:
                self.central_tabs.setCurrentIndex(i)
                return

        doc = self.doc_dao.get_by_id(doc_id)
        if not doc:
            self.show_error("Hata", "Belge bulunamadı.")
            return
        
        doc_text = doc.get('extracted_text') or ""
        if not doc_text.strip():
            show_info(
                self, "Bilgi",
                "Bu belgeler metin içeriği boş. AI sohbeti için belgede metin bulunması gerekiyor."
            )
            return
        
        # Create as a widget for tab integration
        chat_panel = ChatWidget(
            doc_id=doc_id,
            doc_title=doc_title,
            doc_text=doc_text,
            chat_dao=self.chat_dao,
            parent=self
        )
        # Use add_analysis_tab to integrate into central area
        subtitle = f"{len(doc_text):,} karakter yüklendi"
        self.add_analysis_tab(chat_panel, tab_name, subtitle=subtitle)
        self.statusbar.showMessage(f"AI Sohbet başlatıldı: {doc_title}")

    def _on_code_cloud_requested(self, doc_id: int, doc_title: str):
        """Generate and display a Code Cloud for the selected document."""
        # Get all coded segments for this document — filtered by active coder for consistency
        coder_id = getattr(self, 'current_coder_id', None)
        try:
            segments = self.segment_dao.get_by_document(doc_id, coder_id=coder_id)
        except Exception:
            segments = []

        if not segments:
            show_info(self, "Bilgi",
                f"'{doc_title}' belgesinde henüz kodlanmış bölüm bulunmuyor.")
            return

        # Count segment-frequency per code
        from collections import Counter
        code_freq: Counter = Counter()
        code_colors: dict = {}
        for seg in segments:
            name = seg.get("code_name") or "?"
            code_freq[name] += 1
            if name not in code_colors:
                code_colors[name] = seg.get("code_color") or "#4F46E5"

        # Build a word-cloud-style data list but with codes
        freq_list = [(name, count) for name, count in code_freq.most_common(60)]

        # Reuse the word-cloud HTML generator
        try:
            from .visualizations import generate_word_cloud_html
            html_path = generate_word_cloud_html(freq_list, title=f"Kod Bulutu: {doc_title}")
        except Exception:
            # Fallback
            rows = "".join(
                f"<tr><td style='color:{code_colors[n]};font-weight:700;padding:4px 12px'>{n}</td>"
                f"<td style='padding:4px 12px'>{c}</td></tr>"
                for n, c in freq_list
            )
            html_content = f"""<!DOCTYPE html><html><head><meta charset='utf-8'>
            <title>Kod Bulutu</title></head><body style='font-family:sans-serif;padding:20px'>
            <h2>Kod Bulutu: {doc_title}</h2>
            <table border='0' cellspacing='4'><tr><th>Kod</th><th>Segment Sayısı</th></tr>
            {rows}</table></body></html>"""
            with tempfile.NamedTemporaryFile(mode='w', suffix='.html', delete=False, encoding='utf-8') as f:
                f.write(html_content)
                html_path = f.name
            atexit.register(lambda p=html_path: os.path.exists(p) and os.unlink(p))

        # Open in the new standalone visualization dialog
        self._open_visualization(f"Kod Bulutu: {doc_title}", html_path)

    def _on_export_requested(self, doc_id: int, doc_title: str):
        """Export a document as TXT or DOCX."""
        doc = self.doc_dao.get_by_id(doc_id)
        if not doc:
            self.show_error("Hata", "Belge bulunamadı.")
            return

        doc_text = (doc.get("extracted_text") or "").strip()
        if not doc_text:
            show_info(self, "Bilgi", "Bu belgenin metin içeriği boş.")
            return

        safe_title = "".join(c if c.isalnum() or c in " _-" else "_" for c in doc_title)

        # Ask format
        fmt_dlg = QDialog(self)
        fmt_dlg.setWindowTitle("Dışa Aktar")
        fmt_dlg.setFixedSize(320, 160)
        fmt_lay = QVBoxLayout(fmt_dlg)
        fmt_lay.setSpacing(12)
        fmt_lay.setContentsMargins(20, 16, 20, 16)
        fmt_lay.addWidget(QLabel(f"<b>'{doc_title}'</b> için format seçin:"))
        btn_row = QHBoxLayout()

        btn_txt = QPushButton("📄 TXT olarak Kaydet")
        btn_txt.setStyleSheet(
            "QPushButton{background:#F8FAFC;color:#334155;border:1px solid #CBD5E1;"
            "border-radius:8px;padding:10px 14px;font-size:12px;font-weight:600}"
            "QPushButton:hover{background:#F1F5F9}"
        )

        btn_docx = QPushButton("📝 DOCX olarak Kaydet")
        btn_docx.setStyleSheet(
            "QPushButton{background:#4F46E5;color:white;border:none;"
            "border-radius:8px;padding:10px 14px;font-size:12px;font-weight:600}"
            "QPushButton:hover{background:#4338CA}"
        )

        btn_txt.clicked.connect(lambda: (setattr(fmt_dlg, '_choice', 'txt'), fmt_dlg.accept()))
        btn_docx.clicked.connect(lambda: (setattr(fmt_dlg, '_choice', 'docx'), fmt_dlg.accept()))
        btn_row.addWidget(btn_txt)
        btn_row.addWidget(btn_docx)
        fmt_lay.addLayout(btn_row)
        fmt_dlg._choice = None

        if not fmt_dlg.exec():
            return

        choice = fmt_dlg._choice

        if choice == 'txt':
            path, _ = QFileDialog.getSaveFileName(
                self, "Belgeyi TXT Olarak Kaydet", f"{safe_title}.txt", "Metin Dosyası (*.txt)"
            )
            if path:
                with open(path, 'w', encoding='utf-8') as f:
                    f.write(f"{doc_title}\n{'='*len(doc_title)}\n\n{doc_text}")
                self.statusbar.showMessage(f"✅ Belge kaydedildi: {path}")

        elif choice == 'docx':
            path, _ = QFileDialog.getSaveFileName(
                self, "Belgeyi DOCX Olarak Kaydet", f"{safe_title}.docx", "Word Belgesi (*.docx)"
            )
            if path:
                try:
                    from docx import Document
                    document = Document()
                    document.add_heading(doc_title, 0)
                    for para_text in doc_text.split('\n'):
                        if para_text.strip():
                            document.add_paragraph(para_text.strip())
                    document.save(path)
                    self.statusbar.showMessage(f"✅ Belge kaydedildi: {path}")
                except ImportError:
                    txt_path = path.replace('.docx', '.txt')
                    with open(txt_path, 'w', encoding='utf-8') as f:
                        f.write(f"{doc_title}\n{'='*len(doc_title)}\n\n{doc_text}")
                    show_info(
                        self, "Bilgi",
                        "python-docx kurulu değil, TXT olarak kaydedildi.\n"
                        f"Konum: {txt_path}\n\n"
                        "DOCX desteği için: pip install python-docx"
                    )
                    self.statusbar.showMessage(f"✅ TXT olarak kaydedildi: {txt_path}")

    def _on_document_imported(self, file_path: str, file_type: str, folder_id: int = None):
        """Handle document import."""
        from processors import extract_text
        
        try:
            # Extract text from file
            extracted_text = extract_text(file_path)
            
            if not extracted_text:
                show_warning(
                    self, "İçe Aktarma Uyarısı",
                    f"Dosyadan metin çıkarılamadı:\n{file_path}"
                )
                return
            
            # Check if document already exists
            title = Path(file_path).stem
            existing = [d for d in self.doc_dao.get_all() if d['file_path'] == file_path]
            if existing:
                show_info(
                    self, "Zaten Mevcut",
                    f"'{title}' belgesi zaten projeye eklenmiş."
                )
                try:
                    self._on_document_selected(existing[0]['id'])
                except Exception as e:
                    logger.debug(f"Mevcut belge seçilemedi: {e}")
                return
            
            # VALIDATE FOLDER ID
            final_folder_id = folder_id
            if folder_id:
                existing_folder = self.folder_dao.get_by_id(folder_id) if hasattr(self.folder_dao, 'get_by_id') else None
                if not existing_folder:
                     all_folders = {f['id']: f for f in self.folder_dao.get_all()}
                     existing_folder = all_folders.get(folder_id)
                
                if not existing_folder:
                    final_folder_id = None
                    self.statusbar.showMessage("Hedef klasör bulunamadı, ana dizine eklendi.")
            
            # Create document
            try:
                doc_id = self.doc_dao.create(title, file_path, file_type, extracted_text, final_folder_id)
            except Exception as e:
                if "FOREIGN KEY" in str(e):
                    doc_id = self.doc_dao.create(title, file_path, file_type, extracted_text, None)
                    self.statusbar.showMessage("Klasör hatası nedeniyle belge ana dizine eklendi.")
                else:
                    raise e
            
            # Refresh document list
            documents = self.doc_dao.get_all()
            folders = self.folder_dao.get_all()
            self.document_tree.populate_tree(documents, folders)
            
            # Show the new document
            self._on_document_selected(doc_id)
            self.set_dirty()
            self.statusbar.showMessage(f"İçe aktarıldı: {title}")
            
        except FileNotFoundError:
            self.show_error("Dosya Bulunamadı", f"Belirtilen dosya yolunda bulunamadı:\n{file_path}")
        except PermissionError:
            self.show_error("Erişim Engellendi", f"Dosyayı okuma yetkiniz yok veya dosya başka bir program tarafından açık:\n{file_path}")
        except Exception as e:
            self.show_error("İçe Aktarma Hatası", "Dosya işlenirken beklenmedik bir hata oluştu.", e)

    def _get_active_folder_id(self) -> int | None:
        """Get the folder ID of the currently selected item in the document tree."""
        try:
            inner_tree = getattr(self.document_tree, 'tree', self.document_tree)
            model = getattr(self.document_tree, 'model', None)
            
            current_idx = inner_tree.currentIndex()
            if not current_idx.isValid():
                return None
            
            item = model.itemFromIndex(current_idx) if model else None
            if item is None:
                return None
            
            item_type = item.data(Qt.ItemDataRole.UserRole)
            if item_type == "folder":
                return item.data(Qt.ItemDataRole.UserRole + 1)
            elif item_type == "document":
                parent = item.parent()
                if parent and parent.data(Qt.ItemDataRole.UserRole) == "folder":
                    return parent.data(Qt.ItemDataRole.UserRole + 1)
        except Exception:
            pass
        return None

    def _on_document_variables_requested(self, doc_id: int, title: str):
        """Handle request to edit variables for a document."""
        from .variable_dialogs import DocumentVariablesDialog
        dialog = DocumentVariablesDialog(
            doc_id, title, self.var_dao, self.var_value_dao, self
        )
        dialog.exec()
        
    def _on_document_activation_changed(self, doc_id: int, is_active: bool):
        """Handle document activation toggle."""
        self._update_retrieved_segments()
